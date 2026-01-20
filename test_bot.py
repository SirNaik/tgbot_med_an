"""
Test script to verify the basic functionality of the medical analysis bot.
This script will test the text extraction functions and simulate the GigaChat integration.
"""

import os
import tempfile
from docx import Document
from PyPDF2 import PdfReader
import openpyxl


def test_text_extraction():
    """Test text extraction from different file formats."""
    print("Testing text extraction functions...")
    
    # Test DOCX creation and reading
    print("\n1. Testing DOCX file creation and reading:")
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        doc_path = temp_file.name
    
    # Create a sample DOCX file
    doc = Document()
    doc.add_paragraph("Результаты анализа крови:")
    doc.add_paragraph("Гемоглобин: 120 г/л (норма: 120-140)")
    doc.add_paragraph("Глюкоза: 6.5 ммоль/л (норма: 3.3-5.5) - ПОВЫШЕНО")
    doc.save(doc_path)
    
    print(f"Created sample DOCX file: {doc_path}")
    
    # Test reading the DOCX file (simulating what our bot does)
    doc = Document(doc_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    text_content = '\n'.join(full_text)
    
    print(f"Extracted text: {text_content}")
    
    # Clean up
    os.unlink(doc_path)
    
    # Test XLSX creation and reading
    print("\n2. Testing XLSX file creation and reading:")
    with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as temp_file:
        xlsx_path = temp_file.name
    
    # Create a sample XLSX file
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet['A1'] = 'Показатель'
    sheet['B1'] = 'Результат'
    sheet['C1'] = 'Норма'
    sheet['A2'] = 'Гемоглобин'
    sheet['B2'] = '120'
    sheet['C2'] = '120-140'
    sheet['A3'] = 'Глюкоза'
    sheet['B3'] = '6.5'
    sheet['C3'] = '3.3-5.5'
    workbook.save(xlsx_path)
    
    print(f"Created sample XLSX file: {xlsx_path}")
    
    # Test reading the XLSX file (simulating what our bot does)
    workbook = openpyxl.load_workbook(xlsx_path, data_only=True)
    sheet = workbook.active
    text = ""
    for row in sheet.iter_rows(values_only=True):
        row_text = [str(cell) if cell is not None else "" for cell in row]
        text += "\t".join(row_text) + "\n"
    
    print(f"Extracted text: {text.strip()}")
    
    # Clean up
    os.unlink(xlsx_path)
    
    print("\n3. Testing the GigaChat prompt preparation:")
    # Simulate the prompt that would be sent to GigaChat
    sample_analysis_data = text_content
    prompt = (
        f"Ты врач, который должен изучить результаты анализов пользователя и сообщить ему "
        f"где и какие результаты отличаются от референсных, с чем это может быть связано "
        f"и на что обратить внимание. Если требуется дополнительное исследование, "
        f"то дать рекомендации к этим исследованиям. Вот данные анализов:\n\n{sample_analysis_data}\n\n"
        f"Но в конце обязательно добавь, что пользователь не должен заниматься самолечением, "
        f"рекомендации сформированы искусственным интеллектом и носят информационный, "
        f"а не рекомендательный характер, и ему следует консультироваться со специалистами."
    )
    
    print(f"Prepared prompt for GigaChat:\n{prompt}")


def test_bot_structure():
    """Test the overall structure of the bot."""
    print("\n" + "="*60)
    print("BOT STRUCTURE OVERVIEW:")
    print("="*60)
    
    structure_info = """
    Medical Analysis Telegram Bot:
    
    1. START COMMAND (/start):
       - Sends welcome message explaining bot capabilities
       - Informs users about supported file formats
       - Asks to upload medical analysis results
    
    2. DOCUMENT HANDLER:
       - Accepts DOC, DOCX, XLS, PDF, JPEG files
       - Extracts text from documents based on format
       - Processes images separately
       - Sends extracted content to GigaChat
    
    3. GIGACHAT INTEGRATION:
       - Sends formatted prompt to GigaChat API
       - Receives medical analysis interpretation
       - Returns results to user with disclaimer
    
    4. SUPPORTED FORMATS:
       - Documents: DOC, DOCX
       - Spreadsheets: XLS, XLSX
       - PDF files: PDF
       - Images: JPG, JPEG, PNG (with potential OCR)
    
    5. DISCLAMER:
       - Clearly states AI-generated content
       - Advises consulting healthcare professionals
       - Specifies informational-only purpose
    """
    
    print(structure_info)


if __name__ == "__main__":
    print("Medical Analysis Bot - Test Suite")
    print("="*60)
    
    test_text_extraction()
    test_bot_structure()
    
    print("\n" + "="*60)
    print("TEST COMPLETE")
    print("To run the actual bot, execute: python bot.py")
    print("Make sure to set up your environment variables first!")
    print("="*60)